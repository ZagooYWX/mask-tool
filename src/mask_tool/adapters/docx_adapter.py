"""Word文档(.docx)适配器"""

from pathlib import Path
from typing import List, Tuple

from docx import Document

from mask_tool.adapters.base import FileAdapter
from mask_tool.models.detection import DetectionResult, DetectionStatus, Location


class DocxAdapter(FileAdapter):
    """Word文档脱敏适配器"""

    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def process(self, input_path: Path, output_dir: Path) -> Path:
        """
        处理Word文档：
        - 遍历所有段落和表格
        - 在run级别进行文本替换（保持格式）
        """
        doc = Document(str(input_path))
        file_name = input_path.stem
        output_path = output_dir / f"{file_name}_masked.docx"

        # 处理段落
        for para_idx, paragraph in enumerate(doc.paragraphs):
            self._process_paragraph(paragraph, str(input_path), paragraph_index=para_idx)

        # 处理表格
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        self._process_paragraph(
                            paragraph, str(input_path),
                            paragraph_index=para_idx,
                            cell_ref=f"R{row_idx + 1}C{cell_idx + 1}",
                        )

        output_dir.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    def _process_paragraph(
        self,
        paragraph,
        file_path: str,
        paragraph_index: int = None,
        cell_ref: str = None,
    ) -> None:
        """处理单个段落中的所有run"""
        full_text = paragraph.text
        if not full_text.strip():
            return

        # 检测
        results = self.detector.detect(full_text, file_path)

        # 为每条结果设置位置信息
        for r in results:
            r.location.paragraph = paragraph_index
            r.location.cell_ref = cell_ref

        # 策略决策
        results = self.policy.apply(results)

        # 筛选需要替换的结果
        to_replace = [
            r for r in results
            if r.status in (DetectionStatus.AUTO_MASK, DetectionStatus.SUGGEST_MASK)
            and r.text in full_text
        ]

        if not to_replace:
            return

        # 构建替换映射：原文 -> Token
        replace_map: dict[str, str] = {}
        for result in to_replace:
            if result.text not in replace_map:
                if self.masker.irreversible:
                    token = "***"
                else:
                    from mask_tool.models.mapping import TokenMapping
                    token = self.masker.token_gen.generate(result.text, result.text_type)
                    self.masker.mappings.append(TokenMapping(
                        token=token,
                        original=result.text,
                        text_type=result.text_type,
                        confidence=result.confidence,
                    ))
                replace_map[result.text] = token

        # 在run级别执行全段落替换
        self._replace_all_in_paragraph(paragraph, replace_map)

    def _replace_all_in_paragraph(
        self,
        paragraph,
        replace_map: dict[str, str],
    ) -> None:
        """
        在段落的所有run中执行替换，处理跨run的文本

        核心思路：
        1. 拼接所有run的文本，记录每个run的起止位置
        2. 在拼接后的完整文本中查找所有需要替换的位置
        3. 根据位置映射回各个run，修改对应run的文本
        """
        if not paragraph.runs or not replace_map:
            return

        runs = paragraph.runs

        # 构建run位置映射：[(run_index, start_in_full, end_in_full), ...]
        run_spans: List[Tuple[int, int, int]] = []
        pos = 0
        for i, run in enumerate(runs):
            start = pos
            end = pos + len(run.text)
            run_spans.append((i, start, end))
            pos = end

        full_text = "".join(run.text for run in runs)

        # 找出所有需要替换的位置区间
        # replacements: [(start, end, replacement_text), ...]
        replacements: List[Tuple[int, int, str]] = []
        for original, token in replace_map.items():
            search_start = 0
            while True:
                idx = full_text.find(original, search_start)
                if idx == -1:
                    break
                replacements.append((idx, idx + len(original), token))
                search_start = idx + len(original)

        if not replacements:
            return

        # 按起始位置排序，处理重叠
        replacements.sort(key=lambda x: x[0])

        # 执行替换：逐run处理
        # 先标记每个位置应该变成什么字符
        new_chars = list(full_text)
        for start, end, token in replacements:
            new_chars[start:end] = list(token)
        new_full = "".join(new_chars)

        # 将新文本按原run的长度比例分配回各run
        # 但更简单可靠的方式：按原run边界切割新文本
        # 因为替换可能改变长度，需要重新计算边界

        # 策略：对每个run，计算它在new_full中对应的新文本
        # 使用位置映射：原始位置 -> 新文本中的位置
        # 简化方案：如果run数量少（<=3），直接把整个段落文本放到第一个run
        # 如果run数量多，用位置映射

        if len(runs) <= 1:
            runs[0].text = new_full
            return

        # 构建原始位置到新位置的映射
        # old_pos -> new_pos
        pos_map: dict[int, int] = {}
        old_pos = 0
        new_pos = 0
        for start, end, token in replacements:
            # replacement之前的字符位置不变
            while old_pos < start:
                pos_map[old_pos] = new_pos
                old_pos += 1
                new_pos += 1
            # replacement区域
            old_pos = end
            new_pos += len(token)
        # replacement之后的字符
        while old_pos < len(full_text):
            pos_map[old_pos] = new_pos
            old_pos += 1
            new_pos += 1

        # 根据映射切割新文本分配给各run
        for run_idx, old_start, old_end in run_spans:
            if old_start in pos_map and old_end in pos_map:
                new_start = pos_map[old_start]
                new_end = pos_map[old_end]
                runs[run_idx].text = new_full[new_start:new_end]
            else:
                # 边界情况：run文本为空或映射缺失
                runs[run_idx].text = ""
